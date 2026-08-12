# -*- coding: utf-8 -*-
"""
生成项目技术方案说明书 Word 文档 (.docx)
《NILM 空调负荷分解 — 代码数据输入输出、用户配置架构、核心算法流程及终态产物全景说明书》
"""

import os
from pathlib import Path
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_color):
    """为单元格设置背景填充色 (#RRGGBB Hex 格式)"""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """为单元格设置内边距"""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3", sz="4", val="single"):
    """设置表格四边与内部灰线边框"""
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:insideV w:val="none"/><w:left w:val="none"/><w:right w:val="none"/></w:tblBorders>')
        tblPr[0].append(borders)

def add_heading_styled(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F) # 深蓝
    elif level == 2:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x2B, 0x54, 0x7E) # 中蓝
    elif level == 3:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x36, 0x64, 0x8B)
    return p

def add_para_styled(doc, text, bold_prefix=None, space_after=6, line_spacing=1.25):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.bold = True
        r_pre.font.name = "Calibri"
        r_pre.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        r_pre.font.size = Pt(10.5)
        r_pre.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, "F5F7FA")
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    # 为单元格左侧添加浅蓝色粗细线背景
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="none"/><w:left w:val="single" w:sz="24" w:space="0" w:color="3B71CA"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>')
    tcPr.append(borders)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text.strip())
    run.font.name = "Consolas"
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x24, 0x29, 0x2F)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def add_styled_table(doc, headers, rows_data, col_widths=None):
    table = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="D0D7DE", sz="4")
    
    # Header row
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.name = "Calibri"
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(hdr_cells[i], "2B547E")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=100, right=100)
    
    # Data rows
    for r_idx, r_data in enumerate(rows_data):
        row_cells = table.rows[r_idx + 1].cells
        fill_color = "FFFFFF" if r_idx % 2 == 0 else "F7F9FB"
        for c_idx, val in enumerate(r_data):
            row_cells[c_idx].text = ""
            p = row_cells[c_idx].paragraphs[0]
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.name = "Calibri"
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
            set_cell_background(row_cells[c_idx], fill_color)
            set_cell_margins(row_cells[c_idx], top=80, bottom=80, left=100, right=100)
            
    if col_widths:
        for row in table.rows:
            for idx, w in enumerate(col_widths):
                row.cells[idx].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def create_document():
    doc = docx.Document()
    
    # 设置页边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # 文档大标题
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(24)
    p_title.paragraph_format.space_after = Pt(8)
    run_t = p_title.add_run("项目技术方案说明书：\n代码数据输入输出、用户配置架构、核心算法流程及终态产物全景说明书")
    run_t.bold = True
    run_t.font.name = "Calibri"
    run_t.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run_t.font.size = Pt(22)
    run_t.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(24)
    run_s = p_sub.add_run("—— 适用版本：v6.12.6+v6.15.0 / v13+ / v14 扩展模块（2026-08-12） ——")
    run_s.font.name = "Calibri"
    run_s.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run_s.font.size = Pt(11)
    run_s.font.color.rgb = RGBColor(0x55, 0x66, 0x77)
    
    # =========================================================================
    # 第一章：项目代码数据输入输出及用户数据配置整体架构
    # =========================================================================
    add_heading_styled(doc, "第 1 章  项目代码数据输入输出及用户数据配置整体架构", level=1)
    
    add_heading_styled(doc, "1.1 体系架构概览与二层解耦控制模型", level=2)
    add_para_styled(doc, "NILM 空调负荷分解项目采用“批量编排调度层 (run_batch_users.py)”与“单用户任务执行流水线层 (run_user_pipeline.py)”二层分离解耦的核心架构。批量调度层负责对文件系统中的被管理对象进行扫描与设备绑定，加载并解析外部统一指定的 JSON 配置文件，并维护断点续跑原子状态表；流水线层则专注单一执行对象的数据采样勘察、对齐重采样、模型分类回归建模、后处理残差学习、漂移审计与每日多维度监控输出。", bold_prefix="【解耦架构宗旨】：")
    
    arch_chart = """
+-----------------------------------------------------------------------------------------------+
|                             批量调度层 (scripts/run_batch_users.py)                           |
|  - 扫描数据根目录 (data/) 并解析受管理对象 (User/Device) 及其绑定输入源文件链                   |
|  - 统一读取 --time-filter-config JSON 文件，匹配用户定制优先权与兜底层设定                         |
|  - 维护 9 列原子写入状态表 (batch_execution_state.csv)，支持 --resume 断点续跑与异常重试           |
+-----------------------------------------------------------------------------------------------+
                                               |
         (通过序列化 JSON CLI 参数 --time-filter-spec 及环境变量 NILM_USER_* 进行环境封堵与透传)
                                               v
+-----------------------------------------------------------------------------------------------+
|                           单任务流水线引擎 (scripts/run_user_pipeline.py)                      |
|  - 垃圾自动回收，严格执行 _CLEANUP_WHITELIST 顶层隔离白名单契约，不误删调度状态表                |
|  - 串行调度基础数据健康分析、时间网格对齐与算法构建全链路脚本：                                    |
|       [Step 01] scripts/01_audit.py          => 采样频率与数据健康度勘察                       |
|       [Step 02] scripts/02_align_and_feat.py => 15分钟重采样对齐 + 气象/漂移特征工程             |
|       [Step 03] scripts/03_train.py          => 两阶段 GBDT 分解 + 季节 MoE + d87 守卫模型训练   |
|       [Step 04] scripts/04_evaluate.py       => 离线全量评估 + L4 残差校正学习 + 分集预测写盘    |
|       [Step 05] scripts/05_inference.py      => 独立生产推理 + L2 漂移监控 + 质量拆分与统计      |
|       [Step 06] scripts/analyze_on_periods.py => 训前/推前事件行为段级与日级深度分析           |
+-----------------------------------------------------------------------------------------------+
"""
    add_code_block(doc, arch_chart)
    
    add_heading_styled(doc, "1.2 数据输入目录布局与正规匹配语法契约", level=2)
    add_para_styled(doc, "为适应自动化测试、批量运行及生产存量数据处理，框架提供对两种源数据物理布局模式的无缝支持：", bold_prefix="【目录组织模式】：")
    add_para_styled(doc, "（1）规范化双分层目录布局 (v9 标准推荐)：训练数据集与推理数据集做严格物理隔离，分别归放于 data/trains/<device_id>_<user_id>/ 与 data/infers/<device_id>_<user_id>/。该模式彻底防止训练集与推理盲测集发生交叉污染。\n（2）平铺正向匹配布局 (向后兼容)：所有 CSV 均置于 data/ 根目录，批量层通过正规表达式解析文件名自动完成设备、用户及信道关联。")
    
    add_para_styled(doc, "项目针对输入文件制定了标准正规表达式 (Regex) 语法契约，核心输入接口定义如下：", bold_prefix="【文件名正则匹配契约】：")
    
    regex_table_headers = ["文件分类", "适用场景", "文件命名正则表达式匹配语法 (Regex)", "必备表字段列及基本约束"]
    regex_table_data = [
        ["主特征/总线 CSV\n(RE_BUS)", "所有总负荷/多信道主测量序列", r"^e241_(?P<device>[^_]+)_(?P<user>[^-]+)-Ch(?P<ch>\d+)-(?P<start>\d{6})-(?P<end>\d{6})(?P<suffix>(-1|-infer)?)\.csv$", "必备时间戳主键(event_time/time)及连续特征列(load_iden_data*)；支持混合时间字符串解析"],
        ["目标/真值 CSV\n(RE_BR)", "所有真实监测分路及评估对比目标", r"^(?P<user>[^-]+)-(?P<start>\d{6})-(?P<end>\d{6})(?P<suffix>(-1|-infer)?)\.csv$", "必备时间戳主键(time)与分离分路功率列(p1, p2, ...，单位：W)"],
        ["中间多路合并表", "跨时间段文件归并过渡产物", "data/merged_bus.csv, data/merged_branch.csv\ndata/infer_bus.csv, data/infer_branch.csv", "自动由 02 脚本对齐聚合生成，支持训练与推理通道物理独立分离"]
    ]
    add_styled_table(doc, regex_table_headers, regex_table_data, col_widths=[1.2, 1.2, 2.3, 1.8])
    
    add_para_styled(doc, "当用户的业务辨识目标是由多个子分路叠加聚合形成时（如同一空调对应主机+辅热或多机套组），系统无需在表里手工预生成求和列，直接通过配置复合表达式字符串完成即时累加物化：\n"
                         "• 配置语法：\"pA+pB[+pC...]\"（例如 \"p1+p2\" 或 \"p1+p2+p3\"）。\n"
                         "• 按需求和与 NaN 传播防呆机制：数据引擎加载分路表时自动物化一列，列名即为表达式文本（如 \"p1+p2\"）；数值计算严格遵循 skipna=False 语义，即任意一列在某时刻缺失 (NaN)，累加结果直接传播为 NaN。这就从底线规避了“缺采集点被静默设为 0W，致使功耗低估与 F1 严重虚高”的隐患。\n"
                         "• 规范去噪与拦截：支持忽略大小写与首尾空格 (P1 + p2 -> p1+p2)，如出现重复分量 (p1+p1) 直接抛错拒绝。", bold_prefix="【通用复合列目标物化语义 (v13.16)】：")
    
    add_heading_styled(doc, "1.3 集中式用户配置体系 (time_filter_config) 与三级优先权", level=2)
    add_para_styled(doc, "通过 CLI 参数 --time-filter-config <path_to_json> 引入外部统一配置文件。解析模块 (time_filter_utils.py) 遵守三重层级优先权覆写准则：", bold_prefix="【配置层级优先级】：")
    add_para_styled(doc, "1. config[\"<user_id>\"][\"<key>\"] —— 最高优先级：为具体指定的用户 ID 个性化定制的业务选项。\n"
                         "2. config[\"_default\"][\"<key>\"] —— 次高优先级：针对全体未明确列出用户的全局默认模版预设。\n"
                         "3. 代码层内置的静态系统常数兜底值 —— 最低优先级：保底运行，保障容错。")
    add_para_styled(doc, "说明：任意以下划线开头的顶级键名（如 _note_, _comment_, _default 等）会被解析模块判定为框架注释或元说明，自动保护并剔除，不会被当做受控制对象加载。")
    
    add_para_styled(doc, "在单独用户对象字典内部，支持定义六大核心通用配置模块，完整功能定义说明见下表：", bold_prefix="【六大核心配置模块功能说明】：")
    
    config_table_headers = ["配置模块名称", "JSON 参数名称与数据结构", "合法数据取值范围", "核心功能语义与异常处理降级策略"]
    config_table_data = [
        ["(1) 顶层训/推时段过滤\n(v12 增强)", "train / infer 字典，含 include 与 exclude 数组", "双元素闭区间 [start, end]\n短日期 YYYY-MM-DD 或带秒 HH:MM:SS", "先求 include 白名单并集（未配置代表取集全貌），再执行 exclude 黑名单区间查杀剔除。短日期自动扩为当日 00:00:00~23:59:59"],
        ["(2) 辨识分路列名\n(v13.4 / v13.16)", "target_col", "单分路 \"pN\" (N≥0 整数)\n或复合分路 \"pA+pB[+pC...]\"", "指定所建模预测的功率真值通道。优先读本项；为空则试探文件名信道匹配或首个 p* 列，最后回退 \"p1\""],
        ["(3) 启动尖峰自适应守卫\n(v13.1 / v13.5)", "guard_enabled", "bool (true/false) 或 null(未配置)", "true=强制开启；false=强制关停(变频小容量适用)。未设置时系统检查训数据，若 |d87|.max<50W 或发生覆盖率<30%，自动降级关停"],
        ["(4) 分集独立过滤\n(v13.2 细粒度)", "splits.train / splits.val / splits.test", "内部各含 include 和 exclude 数组", "四步一致性处理：(a)初始划分 -> (b)include 硬锚定(优先权 train->val->test) -> (c)跨集均衡让出保持原始切分形状不变 -> (d)exclude 剔除移入重分配池"],
        ["(5) 系统公共常量覆盖\n(v13.5 / v13.13)", "common_overrides 字典，\n支持 9 个系统核心字段", "on_thr_w: [0.001, 5000.0]\nsplit_ratios: [f,f,f] 和为1\nsplit_strategy: 4大切分模式\n纬度:-90~90/经度:-180~180", "提供按用户层覆盖内置常数的通路。推荐切分选 \"global_stratified\"，解决跨月按月硬分层中小月样本强挤往 train 的偏差。参数越界抛 WARN 并回退默认"],
        ["(6) v14 进阶控制选项\n(v14 增强)", "v14_flags 字典或布尔选项", "v14_enable, physics, calibrate, health, diag 等", "向底层流程脚本扩展开启特征后校验、残差自调正层、数据健康诊断表等扩展特性的使能开关"]
    ]
    add_styled_table(doc, config_table_headers, config_table_data, col_widths=[1.3, 1.4, 1.6, 2.2])

    # =========================================================================
    # 第二章：核心算法流程（包括各个模块输入输出）
    # =========================================================================
    add_heading_styled(doc, "第 2 章  核心算法流程与各个模块输入输出分析", level=1)
    add_para_styled(doc, "业务全链条自动化执行由 01 至 06 号标准工作流程脚本与单进程流水线编排模块构成，具备清晰的数据转换关系与环境隔离设计。以下剖析各核心模块的算法原理、输入与输出依赖。", bold_prefix="【全链路模块分解】：")
    
    add_heading_styled(doc, "2.1 [Step 01] 原始数据勘察与质量健康审计模块 (scripts/01_audit.py)", level=2)
    add_para_styled(doc, "为防止上游传感器或数据库出现网络丢包、时间戳断链与严重缺数，而在下游模型层引发难以调适的“幻觉误差”，必须在流程起点实施输入数据健康审计：\n"
                         "• 模块输入：训练与推理的原始总线 CSV 列表及目标分路 CSV 列表。\n"
                         "• 核心算法流程：\n"
                         "  ① 采样间隔分布统计：解析相邻时间戳差值，检查实际物理步长与标称等时间隔 (15分钟=900秒) 的标准差是否偏离。\n"
                         "  ② 连续断层扫描：定位超过容忍间隔 (>1800秒) 的时间跨度段，标注为时序断开区。\n"
                         "  ③ 零值/空值审计：检查各特征及分路序列中的 NaN、无效极大极小噪声比重。\n"
                         "• 模块输出：审计通过后输出控制台健康报告；如判定连续有效率过低将直接提前报警，避免无效算力消耗。", bold_prefix="【逻辑与处理流】：")
    
    add_heading_styled(doc, "2.2 [Step 02] 时间序列重采样对齐与多特征工程模块 (scripts/02_align_and_feat.py)", level=2)
    add_para_styled(doc, "该模块负责把原始采样不均或时间互有偏移的总线与分路数据统一到一个标准化等距离时序网格，并生成环境与漂移特征：\n"
                         "• 模块输入：原始总线 CSV、原始分路 CSV、时段过滤规则参数 --time-filter-spec 及环境常量覆盖参数。\n"
                         "• 核心算法流程：\n"
                         "  ① 时间戳规整与重采样对齐：以 15 分钟为统一基调，构造标准化等距 DatetimeIndex，对总线及分路进行 Inner-Join / Left-Join 时序对齐。\n"
                         "  ② 时段过滤前置执行 (time_filter_utils.py)：对规整后的数据调用 apply_time_filter()，先保留 include 并集，再剔除 exclude 区间；再调用 apply_per_split_filter() 对各分区应用三区细粒度切剪。\n"
                         "  ③ 复合分路列物化：在目标真值表中按 target_col 请求实时执行“按行求和并进行 NaN 传播防呆”。\n"
                         "  ④ 气象维度提取与缓存 (fetch_weather.py)：调用 Open-Meteo Archive API 请求对应位置经纬度的气温及相对湿度特征，并将结果序列化缓存在 models/<user>/weather_cache/ 下，消除重复网络 IO。\n"
                         "  ⑤ L1 漂移感知特征构建 (drift_features.py)：生成近期连续 7 天移动滑动平均电耗等 5 维分布差异特征，通过 7 天暖启动完美消除设备开启早期的温度响应漂移；同时构建 20 桶温度-功率基底表 temp_power_lut.csv。\n"
                         "• 模块输出：标准 15 分钟间距对齐特征-目标矩阵 artifacts/<user>/aligned_15min.csv 及分箱参考 LUT。", bold_prefix="【逻辑与处理流】：")

    add_heading_styled(doc, "2.3 [Step 03] 两阶段 GBDT 建模与自适应守卫模型训练模块 (03_train.py & 14_train_v14.py)", level=2)
    add_para_styled(doc, "本模块是项目预测模型学习的核心构建模块，集成了季节多模型路由、级联树学习及突发尖峰限制守卫：\n"
                         "• 模块输入：对齐特征矩阵 aligned_15min.csv、环境变量 NILM_USER_* 传递的覆盖参数及自适应进阶项。\n"
                         "• 核心算法流程：\n"
                         "  ① 数据分区切分工程 (split_utils.py)：支持 stratified_day（按月保天分层）、time（按时序切划）以及极度推荐的 global_stratified（跨月全局均分抽样，消除按月分层小月被强制挤压致使样本占比失衡）等策略，并写入白名单清单 split_dates_bundle.json。\n"
                         "  ② 季节多模型路由 (Seasonal MoE)：依据环境气温分布把序列路由分配到 Summer（夏天高温层）、Transition（春秋过渡层）和 Winter（冬天低温层）三大专属专家体系建模。\n"
                         "  ③ 两阶段 GBDT 级联建模 (Two-Stage GBDT)：\n"
                         "     • 阶段一 (分类器)：使用 GradientBoostingClassifier / RF 判定目标在给定时刻是否开启 active state (state_pred 0/1)。\n"
                         "     • 阶段二 (连续回归器)：对分类为激活的样本使用 GradientBoostingRegressor 执行回归估值，对输出执行 Isotonic Regression 保证非负和物理递增；采用 F0.5 优化的激活概率阈值寻找最佳分界点，并使用 post_min_on (最小开启时长) 与 post_fill_short_off (短停顿间歇填充) 滤除尖峰噪声。\n"
                         "  ④ d87 自适应尖峰守卫层 (v6.15)：基于开启启动浪涌特征判定设备，引入 Soft-max 连续平滑、样本量自适应调节与开机概率置信融合。若通过训练集最大浪涌核算 |d87|.max < 50W 或出现浪涌点占比低于 30%，模型发出警告并自动降级关停守卫，保障平稳型变频设备的高召回率。\n"
                         "  ⑤ 对照基线构建 (03b_train_v42_baseline.py)：同步训练一份非参数化对照参考基线（Baseline RF 或传统均值策略）。\n"
                         "• 模块输出：持久化主资产 models/<user>/model_bundle_v6_15_0.pkl、rf_model.pkl；并在 artifacts/trains/<user>/ 下完备输出主模型与基准模型在训练与离线评估分集的预测文件（train_pred.csv、val_pred.csv、train_pred_rf.csv、val_pred_rf.csv）。", bold_prefix="【逻辑与处理流】：")

    add_heading_styled(doc, "2.4 [Step 04] 离线全量预测与加性残差校正学习模块 (scripts/04_evaluate.py)", level=2)
    add_para_styled(doc, "基于已建的模型对所有保留及测试分区执行全景推演，并针对系统性分布偏差进行残差回归校正学习：\n"
                         "• 模块输入：model_bundle_v6_15_0.pkl、对齐切分集 (val / test 子集)。\n"
                         "• 核心算法流程：\n"
                         "  ① L4 残差校正回归学习 (residual_calibrator.py)：读取主模型在验证集上的原始估值残差 (val_pred vs y_true_W)，构造高斯/树残差校正器，学习环境温度等协变量与预测偏差的映射关系。\n"
                         "  ② L5 动态多模型智能切调 (v6.9 L4-aware)：在推导中当监测到显著漂移告警时，系统依照校正增益动态平滑切换模型权重。\n"
                         "  ③ 全集性能指标考核：系统化计算回归对齐指标 (SAE, MAE, RMSE, kWh 误差) 与分类对齐指标 (F1, Recall, Precision)，全方位核验非过拟合表现。\n"
                         "• 模块输出：将残差学习网络更新入持久化模型包 model_bundle_v6_15_0.pkl；并在 artifacts/trains/<user>/ 下输出增强后的 test_pred.csv（在文件右侧额外追加校正列 y_pred_W_main_L4_calib 和 residual_W_main_L4_calib）、test_pred_rf.csv 及评估报告图。", bold_prefix="【逻辑与处理流】：")

    add_heading_styled(doc, "2.5 [Step 05] 独立生产盲测推理与动态漂移检测模块 (scripts/05_inference.py)", level=2)
    add_para_styled(doc, "服务在线/离线未标记或真实生产测试数据的推演、概念漂移巡检以及防止将训练日计入指标的数据泄漏审计：\n"
                         "• 模块输入：独立推理特征表 infer_bus.csv、真值对照表 infer_branch.csv（可选）、主模型包 model_bundle_v6_15_0.pkl 及 --time-filter-spec 规范。\n"
                         "• 核心算法流程：\n"
                         "  ① 在线动态时段及分路规整：根据推理要求剔除黑名单区域并执行特征对齐。\n"
                         "  ② L2 即时概念/协变量漂移感知 (drift_detect.py)：对比待推演特征相对于预存 LUT 20 桶训练参考分布的散度偏移，实时给出环境状态结论：OK / WARN (绝对相对漂移≥15%) / ALERT (漂移≥30%) / NO_DATA。\n"
                         "  ③ 并存三级预测计算：在一个工作流内产出主分类-回归原始推导 (raw)、加性残差补偿校正估值 (L4_calib) 与主生产选择输出 (main final)。\n"
                         "  ④ 数据泄漏与盲测边界动态核算 (metrics_utils.py)：将要评估的所有对齐日期同 split_dates_bundle.json 对对碰。若发现存在误入推理段的旧训练日期，程序发出 WARN 并将评估报表自动分列拆成 inference_leak (泄漏部分) 和 inference_ood (正规 OOD 盲测)，有效制止加权平摊导致的假高精度掩盖真实偏离。\n"
                         "• 模块输出：artifacts/infers/<user>/predictions/inference_result.csv 生产全表、25 列逐日指标表 inference_daily_metrics.csv、多算法统计比照 inference_comparison.csv 与协变量漂移诊断文件。", bold_prefix="【逻辑与处理流】：")

    add_heading_styled(doc, "2.6 [Step 06] 连续行为切段与全天静默监测分析模块 (scripts/analyze_on_periods.py)", level=2)
    add_para_styled(doc, "模型推断好坏不可仅看宏观均值，必须向用户揭示开机连续阶段的周期行为并监督静默日分布：\n"
                         "• 模块输入：目标分路数据（训练集或推理集）、判定阈值 on_thr_w。\n"
                         "• 核心算法流程：\n"
                         "  ① 功率二值化切段寻找：基于开启阈值，以滑动算法抓取每一个状态连续处于 Active 态的连通时间窗。\n"
                         "  ② 段级宏观及微观极限统计：运算连续开启过程的开始时间、终止时间、持续时分 duration_min、均值 mean_w、峰值 peak_w、累计贡献度 energy_kwh 及对应归属 dataset；重点扩展最小峰谷审计字段 min_w (v13.16)，确保满足极值不变量 min_w <= mean_w <= peak_w，帮助审计变频空调在连续运转期是否长期稳定维持在下限档位。\n"
                         "  ③ 全日未启动背景静默日构建 (OFF Day Rows, v13.11)：为避免零触发生效天数从总表中消失，系统针对全天没有一个激发有效段的日子，专门生成一条静默汇总明细 (duration_min=1440)，其极值与平均指标取自整天对时待机背景功率序列，实现对日常背景耗电基准的全程监察。\n"
                         "• 模块输出：阶段工作报表 <stage>_on_periods.csv 及按天统计表 <stage>_on_periods_daily.csv。", bold_prefix="【逻辑与处理流】：")

    # =========================================================================
    # 第三章：最终输出产物详解与表结构全集
    # =========================================================================
    add_heading_styled(doc, "第 3 章  最终输出产物详解与表结构全景字典", level=1)
    add_para_styled(doc, "系统运行所产生的一切成果经过严密隔离控制，被持久化归档入四大层次分类，并在 artifacts/trains/<user>/ 下极其完备地保存了主模型与参考对照基线模型在全部三阶段切分集合上的历史预测详表。", bold_prefix="【产物全栈架构】：")

    add_heading_styled(doc, "3.1 产物目录分类体系全景树（含全部分区切分集预测表规范）", level=2)
    add_para_styled(doc, "项目工作区 (workspace_root) 最终生成并持久留存的文件及目录架构展示如下：")

    tree_str = """
<workspace_root>/
 ├── models/                            # [持久化资产包区] 系统运行不会对该区域执行临时文件回收
 │    └── <user_id>/
 │         ├── model_bundle_v6_15_0.pkl # 主模型核心包：含序列标准化转换器、主回归模型、残差模型等
 │         ├── rf_model.pkl             # 对照参考模型包 (Baseline RF)
 │         ├── split_dates_bundle.json  # 数据划分白名单包：保存分配入 train/val/test 的具体日期序列
 │         └── weather_cache/           # 气象协变量 API 请求转换后的离线序列缓存
 └── artifacts/                         # [运行过程产物区] 受顶层白名单契约保护；按处理链路分层
      ├── batch_execution_state.csv     # (白名单保护) 9 列标准的原子化断点续跑执行进度状态表
      ├── batch_run_summary.csv         # (白名单保护) 批处理多对象总精度 KPI 一局展示表
      ├── summary_metrics_all_users.csv # (白名单保护) 全体任务多维度性能横向比较分析报表
      ├── skipped_users.csv             # (白名单保护) 因配置约束或审计未达要求软跳过的用户名单
      ├── trains/                       # 【训练与离线评估产物分类体系】
      │    └── <user_id>/
      │         ├── train_pred.csv      # ⭐ 主模型在训练分区集 (Train Split) 上的全周期预测与残差明细
      │         ├── val_pred.csv        # ⭐ 主模型在验证分区集 (Val Split) 上的全周期预测与残差明细
      │         ├── test_pred.csv       # ⭐ 主模型在测试分区集 (Test Split) 上的完整预测 (含校正残差列)
      │         ├── train_pred_rf.csv   # ⭐ 参考对照模型 (Baseline RF) 在训练集上的预测比对明细
      │         ├── val_pred_rf.csv     # ⭐ 参考对照模型 (Baseline RF) 在验证集上的预测比对明细
      │         ├── test_pred_rf.csv    # ⭐ 参考对照模型 (Baseline RF) 在测试集上的预测比对明细
      │         ├── train_on_periods.csv       # 训前连续开启事件行为的段级明细报表 (含 min_w 极值)
      │         ├── train_on_periods_daily.csv # 训前事件行为每日汇总表 (含背景基线静默日)
      │         └── temp_power_lut.csv         # 训练侧 20 桶温度-功率外表分箱查询表 (LUT)
      └── infers/                       # 【独立生产盲测推理与漂移监控产物分类体系】
           └── <user_id>/
                ├── predictions/
                │    └── inference_result.csv  # 生产级独立推理全时期三级输出多算法对比总表
                ├── metrics/
                │    ├── inference_metrics.csv       # 推理全体均值评估指标
                │    ├── inference_daily_metrics.csv # [25列标准格式] 逐日追踪与真实原始采样密度审计表
                │    └── inference_comparison.csv    # 主预测与各大基准算法在盲测数据上的指标对照表
                ├── infer_on_periods.csv             # 推理验证前连续行为的段级统计表
                ├── infer_on_periods_daily.csv       # 推理阶段每日行为汇总表
                └── inference_temp_power_actual_vs_expected.csv # 环境协变量分区实测漂移评估告警表
"""
    add_code_block(doc, tree_str)

    add_heading_styled(doc, "3.2 模型层持久化资产详解 (models/<user_id>/)", level=2)
    add_para_styled(doc, "在此层级所保存的 pkl 与 json 构成了一组开箱即可直接调用的自包含推理运行单元：\n"
                         "• model_bundle_v6_15_0.pkl：通过 joblib 持久化的核心字典组件。其中包装有特征缩放转换器 (scaler)、季节路由判别阈值、第一阶段分类 GBDT、第二阶段连续回归 GBDT 及由 residual_calibrator 构造的 L4 残差高斯补偿回归模型。\n"
                         "• rf_model.pkl (或 baseline_model.pkl)：对应训练出来的独立对比模型，便于排查算法结构优势与模型复杂度效益比。\n"
                         "• split_dates_bundle.json：精准以字符串数组格式记录本轮次被指定属于训练集 (train_dates)、验证集 (val_dates) 与测试集 (test_dates) 的自然日期清单，供跨环境调用时对目标切分做零差重现，也是识别数据泄漏问题的关键标尺。\n"
                         "• weather_cache/：本地存储的 Open-Meteo 历史气象特征数据文件，使下一次重训能够做到零外部 API 流量快速重加载。", bold_prefix="【持久化模型资产说明】：")

    add_heading_styled(doc, "3.3 训练与离线全部分切集预测表规范 (train/val/test_pred*.csv)", level=2)
    add_para_styled(doc, "为保证在工程验收中，对所有切分集上的方差和拟合表现具备逐点级别的对齐审查能力，模型管道在离线训练阶段 (03_train.py / 04_evaluate.py) 自动为**主模型 (main)** 和**对照基准模型 (Baseline RF)** 各生产 3 组完整输出表，集中于 artifacts/trains/<user_id>/ 下：", bold_prefix="【6 大切分预测明细表】：")
    add_para_styled(doc, "1. train_pred.csv / val_pred.csv / test_pred.csv：记录主建模逻辑分别在对应训练分区、验证分区与测试分区上的完整时序结果。为了能在测试集上以 BI 表格直观展示校正层收益，test_pred.csv 的最右侧自动附加经过 L4 加性后处理校正后的两列数值 (y_pred_W_main_L4_calib 和 residual_W_main_L4_calib)。\n"
                         "2. train_pred_rf.csv / val_pred_rf.csv / test_pred_rf.csv：在与主模型完全等同的分组下标和样本点下，对照基准算法所输出的结果与残差序列，构成客观、非参数化的参照刻度。")
    
    add_para_styled(doc, "上述 6 大切分集预测明细文件的标准通用列结构定义见下表：", bold_prefix="【切分表通序列结构字典】：")
    
    split_pred_headers = ["字段名", "数据类型", "所属列属性", "业务工程意义与说明"]
    split_pred_data = [
        ["time", "ISO String", "必备列", "重采样规整后的 15 分钟间距时间戳序列 (yyyy-MM-dd HH:MM:SS)"],
        ["y_true_W", "Float", "必备列", "给定输入分路中的原始真实标杆目标功率 (W)，保留 3 位小数"],
        ["y_pred_W", "Float", "必备列", "当前处理模型在该时间节点推导出的预测有效目标值 (W)"],
        ["residual_W", "Float", "必备列", "当前主预测相对客观标准值的绝对差值 (y_pred_W - y_true_W)"],
        ["state_true", "Int (0/1)", "可选/附加列", "按阈值通过客观真值判定出的二值化开关活动事件态 (Active State)"],
        ["state_pred", "Int (0/1)", "可选/附加列", "模型在阶段一分类推演下，判定该目标在此刻是否触发的开关响应"],
        ["p_on", "Float (0~1)", "可选/附加列", "模型向外释放的目标在此时刻处在开启激发态的置信概率度量"],
        ["y_pred_low_W /\ny_pred_high_W", "Float", "可选/附加列", "连续回归区间估算的统计下边界与统计上界预测"],
        ["y_pred_W_main_L4_calib /\nresidual_W_main_L4_calib", "Float", "仅在 test_pred.csv\n中追加呈现", "施加后处理加性高斯校正层 (L4) 补偿调整后的独立预测功率及残差"]
    ]
    add_styled_table(doc, split_pred_headers, split_pred_data, col_widths=[1.5, 1.1, 1.3, 2.6])

    add_heading_styled(doc, "3.4 生产盲测独立推理产物规范 (predictions/inference_result.csv)", level=2)
    add_para_styled(doc, "用于在线或全期盲测推理环节。在此单表中统一呈现了主分类-回归逻辑下的“三级预测口径”与各参考对照算法指标，供直接生产使用和多算法效益横向评估：", bold_prefix="【独立生产盲测输出字典】：")
    
    infer_headers = ["列名标识", "数据类型", "所属算法口径", "核心含义与说明"]
    infer_data = [
        ["time / event_time", "ISO String", "基底主键", "对齐后的 ISO 标准连续统一时序主键"],
        ["y_true", "Float", "实际真值", "真实的客观目标值（若生产场景纯无标签盲测推理则允许缺省空置）"],
        ["y_pred_W_main", "Float", "主模型终态口径", "最终对下游调用的正式预测成果（含两阶段推导 + L4校正 + L5切换）"],
        ["y_pred_W_main_raw", "Float", "主模型原始口径", "未做 L4 补偿且无 L5 加权的原始 GBDT 模型回归输出，提供基础对照"],
        ["y_pred_W_main_L4_calib", "Float", "校正层验证口径", "仅叠加后处理残差补偿层后的结果（不含 L5 加权切分层改变）"],
        ["state_pred_main", "Int (0/1)", "主分类决策", "主模型评估的此点目标是否处于激活运作响应态"],
        ["p_on_main", "Float (0~1)", "概率度量", "主模型给出的时刻响应态概率（配合评估最佳决策截断值 F0.5）"],
        ["residual_W_main_raw /\nresidual_W_main_L4_calib", "Float", "残差列组", "原始推导对真值的残差与残差校正后预测对真值的绝对残差对比"],
        ["y_pred_W_rf ...", "Float", "基准对照口径", "经过同步对齐归并而来的 Baseline RF 等算法在该时刻的参考估值"]
    ]
    add_styled_table(doc, infer_headers, infer_data, col_widths=[1.6, 1.1, 1.3, 2.5])

    add_heading_styled(doc, "3.5 25 列逐日指标表与原始采样密度审计表规范 (daily_metrics.csv)", level=2)
    add_para_styled(doc, "系统在离线阶段与独立推理阶段分别写就 train_daily_metrics.csv 和 inference_daily_metrics.csv。本说明书确立了其中 25 列严格标准表格式；最重要的在于增补了第 5 列 n_bus_raw 和第 6 列 n_branch_raw，能够让维护团队瞬间判断“当天的低 F1 或大误差指标，究竟是源于模型自身的泛化偏差，还是因为真实采集端原 CSV 中大量漏数”：", bold_prefix="【25 列标准评估表字典】：")
    
    daily_headers = ["顺位索引", "列字段名称", "类型", "含义解释与特别审计语义说明"]
    daily_data = [
        ["1", "date", "String", "统计目标自然日 (YYYY-MM-DD)"],
        ["2", "split", "String", "该记录所属数据域 (train / val / test / inference)；若遇到实际推理区间包含已划拨入历史训练集的日期，程序自动按日期将其拆分输出为两行：inference_leak (泄漏统计行) 与 inference_ood (纯盲测 OOD 行)"],
        ["3", "model", "String", "评价对象算法来源名称（主程序为 main，对照组为 rf 等）"],
        ["4", "n_samples", "Int", "此统计日在通过对齐和过滤等距网格后实际具备的计算测试记录条数"],
        ["5", "n_bus_raw\n(⭐审计专用)", "Int / \"\"", "上游主特征源表当天物理读取行数：在此记录了不经对齐的原表实际行数。若该数值显著低于 288 条 (15分钟一整天满载记录数)，可作为不可抗源头漏数凭据"],
        ["6", "n_branch_raw\n(⭐审计专用)", "Int / \"\"", "目标/分路源表当天物理读取行数：提供实际有效采样点条目数，与上面互相印证"],
        ["7 ~ 11", "Accuracy, Precision,\nRecall, F1, AUC", "Float", "分类状态判定综合绩效矩阵：各项均保留 4~6 位精度小数，若缺失则为 \"\""],
        ["12 ~ 14", "MAE_W, RMSE_W,\nSAE", "Float", "回归预测指标：其中 SAE = abs(kWh_pred - kWh_true)/max(kWh_true, 0.01)，直观体现全天的宏观电耗估计累积偏度"],
        ["15 ~ 17", "kWh_true, kWh_pred,\nkWh_err", "Float", "单日累计工作能耗度数分析：真实总度数 / 预估总度数 / 二者差值比对"],
        ["18 ~ 21", "TP, FP, FN, TN", "Int", "事件二值分类判定混淆矩阵基础统计项 counts"],
        ["22", "dataset", "String", "标记本行数据来源归属切片标记 (如 train, val, test, used, excluded)"],
        ["23", "on_thr_w", "Float", "本行指标评估所应用的绝对判定活动阈值瓦数 (W)"]
    ]
    add_styled_table(doc, daily_headers, daily_data, col_widths=[0.8, 1.4, 1.0, 3.3])

    add_heading_styled(doc, "3.6 事件连续段统计表与静默天数表规范 (<stage>_on_periods*.csv)", level=2)
    add_para_styled(doc, "为深度监测连续开启行为和待机能耗规律，系统在训前及推导后生成段级及日级报表：", bold_prefix="【段级及日级统计说明】：")
    add_para_styled(doc, "• <stage>_on_periods.csv (段表)：记载各个联通活动响应期的统计全貌。核心列包含 being_time (开始点)、end_time (终结点)、p1 (目标通道)、duration_min (持续分长)、mean_w (时段平均功率)、peak_w (峰顶数值)、energy_kwh (此段耗电度数)、dataset (切块归属)。最关键扩展在于第 5 列 min_w (连续段最小瞬时极值)：可清晰观察此段最低落入的功率边界，且确保满足恒等约束 min_w <= mean_w <= peak_w。当出现一整天没有触发动作的日子，程序自动补充一条全日静默记录 (OFF Day Row，duration_min=1440)，保证非活动日背景待机分布不漏报。\n"
                         "• <stage>_on_periods_daily.csv (日表)：汇成按日宏观行为卡片。包含 date、n_segments (日激发频度，静默日=0)、total_on_hours、first_on_time、last_off_time 及均值积分统计。")

    add_heading_styled(doc, "3.7 协变量参考基底表与实测漂移评估规范 (lut.csv & actual_vs_expected.csv)", level=2)
    add_para_styled(doc, "• 训练标准桶表 (temp_power_lut.csv)：系统按外源气温划分 20 个等跨度基本桶 (Bin 1~20)，结合 1 个全样本兜底中位桶 ALL_MEDIAN (共 21 行)。每行包含分箱上下界 temp_lo/temp_hi、参考目标中位数 expected_signal、样本计数 n_samples 及均值与标准差。\n"
                         "• 推导漂移比照表 (inference_temp_power_actual_vs_expected.csv)：利用推理样本向对应 20 桶填注实测中位数，经两端核查算得 abs_residual 与相对漂移率 rel_drift；随后生成至关重要的状态评估信号列 drift_flag，明确给出四项状态结论：\n"
                         "  ① OK —— 推理区间指标未明显偏离基准，模型处在稳态；\n"
                         "  ② WARN —— 测点产生漂移警报（绝对偏差率超范围，如 ≥15%）；\n"
                         "  ③ ALERT —— 检测到严重变异风险（偏移发生巨大恶化，如 ≥30%），触发对下游的降权与限制措施；\n"
                         "  ④ NO_DATA —— 此区间内当前批次没有采集点，安全空跳。", bold_prefix="【协变量分区基底表字典】：")

    add_heading_styled(doc, "3.8 批量原子执行状态表规范 (artifacts/batch_execution_state.csv)", level=2)
    add_para_styled(doc, "作为面向成百上千对象的批量调度中枢资产，断点续跑管理不仅要有精准的状态记录，更要具备抗断电、抗系统异常宕机的强容灾操作保护：", bold_prefix="【9 列原子断点状态表协议】：")
    add_para_styled(doc, "• 无缝原子写就双重保险：框架处理本表更新时，严令制止修改现成主文件。程序采取“先在内存生成数据后完整写盘生成 batch_execution_state.csv.tmp 临时文件 ➔ 写就后调取系统底层 atomic rename 接口 (os.replace) 无缝覆盖原表”的保护链路。若运行期间任意阶段进程发生被动异常中斩，盘内旧表内容仍百分百合法一致，最高仅会丢失正在计算阶段的当下一行。\n"
                         "• 表结构定义 (UTF-8-BOM)：完全固化为 9 个标准列表头规范。\n"
                         "  1. user_id: 正在执行管理的对象主键或账户代号；\n"
                         "  2. status: 计算完结状态表达 (ok / fail / soft_skip)；\n"
                         "  3. success: Bool 布尔标记指示 (True / False)；\n"
                         "  4. started_at / finished_at: 任务启动与交盘时间戳；\n"
                         "  5. duration_s: 本轮次任务全量消耗分秒 Float；\n"
                         "  6. message: 执行提示信息文字摘要，支持中文反馈错误原因；\n"
                         "  7. target_col: 经计算并进行累加判定后的目标分路名称 (如 p1+p2)；\n"
                         "  8. run_id: 把本批状态行绑定到当前的调起批次码。")

    # 保存文件
    out_path = Path("/home/user/nilm_test/项目技术方案说明书_数据架构与核心算法全景规范.docx")
    doc.save(str(out_path))
    print(f"成功创建技术方案说明 Word 文档：{out_path}，大小：{out_path.stat().st_size} 字节")

if __name__ == "__main__":
    create_document()
